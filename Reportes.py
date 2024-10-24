import streamlit as st
import requests
import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
from docx import Document
from docx.shared import Inches
from datetime import datetime
from io import BytesIO

# Título de la aplicación
st.title("Generación automática de reportes de la TRM")

# Colocamos la URL de la API
url = "https://www.datos.gov.co/resource/ceyp-9c7c.json"

# Inicializamos el DataFrame vacío
df = pd.DataFrame()

# Colocamos los parámetros de la extracción
limit = 10000
offset = 0

while True:
    params = {
        "$limit": limit,
        "$offset": offset
    }

    # Hacemos la solicitud GET a la API con los parámetros de la extracción
    response = requests.get(url, params=params)
    
    # Verificamos si la solicitud fue exitosa
    if response.status_code == 200:
        # Convertimos la respuesta en JSON
        data = response.json()
        
        # Convertimos los datos a un DataFrame temporal
        temp_df = pd.DataFrame(data)
        
        if temp_df.empty:
            break
        
        # Concatenamos los datos al DataFrame principal
        df = pd.concat([df, temp_df], ignore_index=True)
        
        offset += limit
    else:
        st.error(f"Error al obtener los datos: {response.status_code}")
        break

# Convertimos las columnas de fecha a datetime
df['vigenciadesde'] = pd.to_datetime(df['vigenciadesde'])
df['vigenciahasta'] = pd.to_datetime(df['vigenciahasta'])

# Formateamos las fechas en el formato 'YYYY-MM-DD'
df['vigenciadesde'] = df['vigenciadesde'].dt.strftime('%Y-%m-%d')
df['vigenciahasta'] = df['vigenciahasta'].dt.strftime('%Y-%m-%d')

# Extendemos el dataframe
df_extendido = df.copy()
new_rows = []
for index, row in df_extendido.iterrows():
    vigenciadesde = pd.to_datetime(row['vigenciadesde'])
    vigenciahasta = pd.to_datetime(row['vigenciahasta'])
    
    if vigenciahasta > vigenciadesde:
        for day in pd.date_range(start=vigenciadesde, end=vigenciahasta):
            new_row = row.to_dict()
            new_row['vigenciadesde'] = day
            new_row['vigenciahasta'] = day
            new_rows.append(new_row)
    else:
        new_rows.append(row.to_dict())

df_extendido = pd.DataFrame(new_rows)

# Renombramos y eliminamos columnas
df_extendido = df_extendido.rename(columns={'vigenciadesde': 'fecha'})
df_extendido = df_extendido.drop('vigenciahasta', axis=1)

# Convertimos la columna 'valor' a tipo float
df_extendido['valor'] = df_extendido['valor'].astype(float)

# Agrupamos por año
df_extendido['fecha'] = pd.to_datetime(df_extendido['fecha'])
df_extendido['year'] = df_extendido['fecha'].dt.year
grouped = df_extendido.groupby('year')

# Preparamos el gráfico
plt.figure(figsize=(10, 6))
for year, group in grouped:
    plt.scatter(group['fecha'], group['valor'], label=str(year), color='g', marker='.')

plt.xlabel('Fecha')
plt.ylabel('Valor')
plt.title('Diagrama de Dispersión Valor vs Fecha (Anual)')
plt.grid(True)
plt.tight_layout()

# Guardamos el gráfico en un archivo temporal
buffer = BytesIO()
plt.savefig(buffer, format="png")
buffer.seek(0)

# Mostramos el gráfico en la app de Streamlit
st.image(buffer, caption="Gráfico de dispersión TRM", use_column_width=True)

# Botón para generar y descargar el reporte
if st.button("Generar y descargar reporte"):
    # Creamos el documento Word
    doc = Document()
    doc.add_heading('Reporte de la TRM', level=1)
    
    # Hallamos los valores de interés
    valor_actual = df_extendido['valor'].iloc[-1]
    max_valor = df_extendido['valor'].max()
    min_valor = df_extendido['valor'].min()
    promedio_valor = df_extendido['valor'].mean()

    # Añadimos al documento
    doc.add_heading('Valores relevantes:', level=2)
    doc.add_paragraph(f"El valor máximo alcanzado es: {max_valor}")
    doc.add_paragraph(f"El valor mínimo alcanzado es: {min_valor}")
    doc.add_paragraph(f"El valor promedio es: {promedio_valor}")
    doc.add_paragraph(f"El último valor es: {valor_actual}")

    # Agregamos la gráfica
    doc.add_heading('Gráfico de TRM:', level=2)
    doc.add_paragraph('Diagrama de dispersión Valor vs Fecha (Anual):')
    
    # Insertamos la imagen en el documento
    doc.add_picture(buffer, width=Inches(4))

    # Guardamos el documento en memoria
    byte_io = BytesIO()
    doc.save(byte_io)
    byte_io.seek(0)

    # Descargar el archivo
    st.download_button(
        label="Descargar informe",
        data=byte_io,
        file_name=f"TRM_Reporte_{datetime.now().strftime('%Y-%m-%d')}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
